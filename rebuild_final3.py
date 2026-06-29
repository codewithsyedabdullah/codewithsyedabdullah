"""v3 with v2 spacing, separators at 80exact, thin spacers removed"""
from docx import Document
from lxml import etree

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
doc = Document(r'C:\Users\786 COMPUTERS\Downloads\resume\Resume_Syed_Abdullah_Yaqoob_v3.docx')

def is_empty_para(p_elem):
    for t in p_elem.iter(f'{{{W_NS}}}t'):
        if t.text and t.text.strip():
            return False
    return True

def set_spacing(p_elem, line, rule, after='0'):
    ppr = p_elem.find(f'{{{W_NS}}}pPr', None)
    if ppr is None:
        ppr = etree.SubElement(p_elem, f'{{{W_NS}}}pPr')
    spacing = ppr.find(f'{{{W_NS}}}spacing', None)
    if spacing is None:
        spacing = etree.SubElement(ppr, f'{{{W_NS}}}spacing')
    spacing.set(f'{{{W_NS}}}line', str(line))
    spacing.set(f'{{{W_NS}}}lineRule', rule)
    spacing.set(f'{{{W_NS}}}after', after)

for p in doc.paragraphs:
    t = p.text
    if t == '':
        # Major section separator - v2 style 4pt
        set_spacing(p._element, 80, 'exact')
    elif t is not None and t.strip() == '' and len(t) <= 2:
        # Thin spacer between subsections - nearly 0
        set_spacing(p._element, 1, 'exact')
    else:
        # Body - v2 style
        set_spacing(p._element, 200, 'auto')

out = r'C:\Users\786 COMPUTERS\Downloads\resume\Resume_Syed_Abdullah_Yaqoob_v3_fixed.docx'
doc.save(out)
print(f'Saved: {out}')

# Quick verify
doc2 = Document(out)
for i, p in enumerate(doc2.paragraphs):
    ppr = p._element.find(f'{{{W_NS}}}pPr', None)
    if ppr is not None:
        spacing = ppr.find(f'{{{W_NS}}}spacing', None)
        if spacing is not None:
            line = spacing.get(f'{{{W_NS}}}line')
            rule = spacing.get(f'{{{W_NS}}}lineRule')
            if line in ('80', '1'):
                txt = p.text[:40] if p.text else '(empty)'
                print(f'P{i}: line={line} rule={rule} {repr(txt)}')
