"""Rebuild v3: v2 spacing (line=200 body), minimize spacers for 1-page fit"""
from docx import Document
from lxml import etree

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
doc = Document(r'C:\Users\786 COMPUTERS\Downloads\resume\Resume_Syed_Abdullah_Yaqoob_v3.docx')

def is_empty_para(p_elem):
    for t in p_elem.iter(f'{{{W_NS}}}t'):
        if t.text and t.text.strip():
            return False
    return True

def is_thin_spacer(p):
    t = p.text
    return t is not None and t.strip() == '' and len(t) <= 2

def set_spacing(p_elem, line, rule, after='0'):
    ppr = p_elem.find(f'{{{W_NS}}}pPr', None)
    if ppr is None:
        ppr = etree.SubElement(p_elem, f'{{{W_NS}}}pPr')
    spacing = ppr.find(f'{{{W_NS}}}spacing', None)
    if spacing is None:
        spacing = etree.SubElement(ppr, f'{{{W_NS}}}spacing')
    spacing.set(f'{{{W_NS}}}line', str(line))
    spacing.set(f'{{{W_NS}}}lineRule', rule)
    spacing.set(f'{{{W_NS}}}after', after)

for i, p in enumerate(doc.paragraphs):
    if is_empty_para(p._element):
        # Section separator - 2pt gap
        set_spacing(p._element, 40, 'exact')
    elif is_thin_spacer(p):
        # Thin space between subsections - small gap
        set_spacing(p._element, 60, 'exact')  # ~3pt visible gap
    else:
        # Body text - exactly v2 style
        set_spacing(p._element, 200, 'auto')

out = r'C:\Users\786 COMPUTERS\Downloads\resume\Resume_Syed_Abdullah_Yaqoob_v3_fixed.docx'
doc.save(out)
print(f'Saved: {out}')

# Verify spacing
doc2 = Document(out)
non_standard = []
for i, p in enumerate(doc2.paragraphs):
    ppr = p._element.find(f'{{{W_NS}}}pPr', None)
    if ppr is not None:
        spacing = ppr.find(f'{{{W_NS}}}spacing', None)
        if spacing is not None:
            line = spacing.get(f'{{{W_NS}}}line')
            rule = spacing.get(f'{{{W_NS}}}lineRule')
            if line != '200' or rule != 'auto':
                txt = p.text[:50] if p.text else '(empty)'
                non_standard.append(f'P{i}: line={line} rule={rule} {repr(txt)}')

print(f'\nNon-standard spacing ({len(non_standard)}):')
for s in non_standard:
    print(f'  {s}')
