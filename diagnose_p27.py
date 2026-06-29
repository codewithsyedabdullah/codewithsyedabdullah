"""Diagnose original P27 text structure"""
from docx import Document
from lxml import etree

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
doc = Document(r'C:\Users\786 COMPUTERS\Downloads\resume\Resume_Syed_Abdullah_Yaqoob_v3.docx')

# P27 in original doc
p27 = doc.paragraphs[27]
elem = p27._element

# Full text
full = ''
for t in elem.iter(f'{{{W_NS}}}t'):
    full += t.text or ''

print(f'=== P27 (index 27) ===')
print(f'Full text length: {len(full)}')
print(f'Repr: {repr(full[:300])}')
print()

# Check for w:br elements
brs = elem.findall(f'.//{{{W_NS}}}br')
print(f'w:br elements: {len(brs)}')

# Check for newlines in text
import re
newline_positions = [m.start() for m in re.finditer(r'\n', full)]
print(f'Newline positions: {newline_positions}')

# Check all t elements and their text
ts = list(elem.iter(f'{{{W_NS}}}t'))
print(f'Total t elements: {len(ts)}')
for i, t in enumerate(ts):
    txt = t.text or ''
    print(f'  t[{i}]: len={len(txt)}, text={repr(txt[:100])}')

print()
# Also check surrounding paragraphs
for offset in range(-2, 3):
    idx = 27 + offset
    if 0 <= idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        t = p.text[:80] if p.text else '(empty)'
        print(f'P{idx}: {repr(t)}')
