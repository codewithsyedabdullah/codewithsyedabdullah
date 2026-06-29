import docx
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx.enum.text

doc = docx.Document(r"C:\Users\786 COMPUTERS\Downloads\resume\Resume_Syed_Abdullah_Yaqoob_v3.docx")

align_map = {
    WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
    WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
    WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
    None: "None (inherited/undefined)",
}

line_spacing_rule_map = {
    docx.enum.text.WD_LINE_SPACING.SINGLE: "SINGLE",
    docx.enum.text.WD_LINE_SPACING.ONE_POINT_FIVE: "ONE_POINT_FIVE",
    docx.enum.text.WD_LINE_SPACING.DOUBLE: "DOUBLE",
    docx.enum.text.WD_LINE_SPACING.AT_LEAST: "AT_LEAST",
    docx.enum.text.WD_LINE_SPACING.EXACTLY: "EXACTLY",
    docx.enum.text.WD_LINE_SPACING.MULTIPLE: "MULTIPLE",
    None: "None",
}

def fmt_val(v):
    if v is None:
        return "None"
    if isinstance(v, Pt):
        return f"{v.pt}pt"
    if isinstance(v, Inches):
        return f"{v.inches}in"
    if isinstance(v, Emu):
        return str(v)
    return str(v)

print("=" * 80)
print("PARAGRAPHS")
print("=" * 80)

for i, para in enumerate(doc.paragraphs):
    print(f"\n{'-' * 70}")
    print(f"PARAGRAPH {i}:")
    print(f"  Text: {para.text!r}")
    print(f"  Style name: {para.style.name if para.style else 'None'}")

    pf = para.paragraph_format
    print(f"  -- Paragraph Format --")
    print(f"    alignment: {align_map.get(pf.alignment, str(pf.alignment))}")
    print(f"    space_before: {fmt_val(pf.space_before)}")
    print(f"    space_after: {fmt_val(pf.space_after)}")
    print(f"    line_spacing: {pf.line_spacing}")
    print(f"    line_spacing_rule: {line_spacing_rule_map.get(pf.line_spacing_rule, str(pf.line_spacing_rule))}")

    if para.runs:
        for j, run in enumerate(para.runs):
            rf = run.font
            print(f"\n  -- Run {j} --")
            print(f"    Text: {run.text!r}")
            print(f"    font.name: {rf.name}")
            print(f"    font.size: {fmt_val(rf.size)}")
            print(f"    font.bold: {rf.bold}")
            print(f"    font.italic: {rf.italic}")
            print(f"    font.underline: {rf.underline}")
            print(f"    font.color.rgb: {rf.color.rgb if rf.color and rf.color.rgb else 'None'}")
            print(f"    font.color.theme_color: {rf.color.theme_color if rf.color else 'None'}")
    else:
        print(f"  (No runs)")

if doc.tables:
    print(f"\n\n{'=' * 80}")
    print("TABLES")
    print('=' * 80)
    for ti, table in enumerate(doc.tables):
        print(f"\n{'-' * 70}")
        print(f"TABLE {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                text = cell.text[:300] if cell.text else ""
                print(f"  Cell[{ri}][{ci}]: {text!r}")
else:
    print(f"\n\n(No tables found)")
