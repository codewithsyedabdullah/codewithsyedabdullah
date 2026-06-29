"""v3 with v2 spacing, bullet splitting at clean boundaries, ALL words preserved"""
from docx import Document
from lxml import etree
import copy

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
doc = Document(r'C:\Users\786 COMPUTERS\Downloads\resume\Resume_Syed_Abdullah_Yaqoob_v3.docx')

def get_all_runs_ordered(p_elem):
    runs = list(p_elem.iter(f'{{{W_NS}}}r'))
    return runs

def p_text(p):
    parts = []
    for t in p._element.iter(f'{{{W_NS}}}t'):
        if t.text:
            parts.append(t.text)
    return ''.join(parts)

def find_first_run(p_elem):
    """Find first w:r element in document order (including inside hyperlinks)"""
    return p_elem.find(f'{{{W_NS}}}r', None) or p_elem.find(f'.//{{{W_NS}}}r', None)

def clear_all_text(p_elem):
    for t in p_elem.iter(f'{{{W_NS}}}t'):
        t.text = ''

def set_first_t_text(p_elem, text):
    """Set text on first w:t element"""
    first_t = p_elem.find(f'.//{{{W_NS}}}t', None)
    if first_t is not None:
        first_t.text = text
        first_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

def set_spacing(p_elem, line, rule, after='0'):
    ppr = p_elem.find(f'{{{W_NS}}}pPr', None)
    if ppr is None:
        ppr = etree.SubElement(p_elem, f'{{{W_NS}}}pPr')
    spacing = ppr.find(f'{{{W_NS}}}spacing', None)
    if spacing is None:
        spacing = etree.SubElement(ppr, f'{{{W_NS}}}spacing')
    spacing.set(f'{{{W_NS}}}line', str(line))
    spacing.set(f'{{{W_NS}}}lineRule', rule)
    spacing.set(f'{{{W_NS}}}after', after)

# Step 1: Fix spacing
for p in doc.paragraphs:
    t = p.text
    if t == '':
        set_spacing(p._element, 80, 'exact')
    elif t is not None and t.strip() == '' and len(t) <= 2:
        set_spacing(p._element, 1, 'exact')
    else:
        set_spacing(p._element, 200, 'auto')

# Step 2: Split employment bullets at clean boundaries
# Each split must: preserve ALL words, make each half a valid standalone bullet

# Define splits: (orig_bullet_text, split_index)
# split_index = character position where second part begins

BULLET_SPLITS = {}

def get_bullet_char(text):
    for c in ['\u2022', '\u25CF', '\uFFFD', '\u2023', '\u2219']:
        if text.startswith(c):
            return c
    return ''

# P18: "Boosted AI content ranking accuracy by 35% (A/B test) by engineering LLM pipelines deployed on AWS Lambda and S3 for scalable serverless inference."
# Split at "by" (the second one). Part1 gets the first clause. Part2 gets "By ..."
# 1. "Boosted AI content ranking accuracy by 35% (A/B test)"
# 2. "By engineering LLM pipelines deployed on AWS Lambda and S3 for scalable serverless inference"
P18_TEXT = "Boosted AI content ranking accuracy by 35% (A/B test) by engineering LLM pipelines deployed on AWS Lambda and S3 for scalable serverless inference."
# Find second "by" - it's at position 59 (after "(A/B test) ")
# Actually: "Boosted AI content ranking accuracy by 35% (A/B test) by"
# Let me count: "Boosted AI content ranking accuracy by 35% (A/B test) " = 57 chars
# Position of second " by " = 56

# P19: "Reduced backend API latency by 40% (profiler benchmarks) by building RESTful Node.js/Express microservices on AWS EC2 with Docker containerization."
P19_TEXT = "Reduced backend API latency by 40% (profiler benchmarks) by building RESTful Node.js/Express microservices on AWS EC2 with Docker containerization."

# P20: "Cut deployment time by 50% (sprint log) by automating CI/CD pipelines via GitHub Actions with Kubernetes orchestration for zero-downtime rollouts."
P20_TEXT = "Cut deployment time by 50% (sprint log) by automating CI/CD pipelines via GitHub Actions with Kubernetes orchestration for zero-downtime rollouts."

# P22: "Scaled membership by 40% (sign-up data) by designing a full-year Agile/Scrum roadmap of 6+ events/semester and deploying the society site on AWS S3 + Vercel."
# Split at "and" - two parallel actions
P22_TEXT = "Scaled membership by 40% (sign-up data) by designing a full-year Agile/Scrum roadmap of 6+ events/semester and deploying the society site on AWS S3 + Vercel."

# P23: "Cut event setup time by 30% (ops log) by automating infrastructure via CI/CD GitHub Actions workflows with Salesforce CRM for member engagement tracking."
P23_TEXT = "Cut event setup time by 30% (ops log) by automating infrastructure via CI/CD GitHub Actions workflows with Salesforce CRM for member engagement tracking."

# P25: "Reduced integration defects by 30% across 3 prototypes by developing modular Python/Node.js components with TensorFlow inference and Pandas/NumPy data pipelines."
# Split at "with" - method vs tools
P25_TEXT = "Reduced integration defects by 30% across 3 prototypes by developing modular Python/Node.js components with TensorFlow inference and Pandas/NumPy data pipelines."

# P26: "Accelerated prototype delivery by 25% (sprint velocity) by containerizing services with Docker and orchestrating via Kubernetes on AWS EC2."
# Split at "and" - two parallel actions  
P26_TEXT = "Accelerated prototype delivery by 25% (sprint velocity) by containerizing services with Docker and orchestrating via Kubernetes on AWS EC2."

# P27: "Improved codebase maintainability by 20% (peer review) via RESTful API design, CI/CD pipelines, and technical documentation across all prototypes.\n\nPROJECTS"
P27_TEXT_BEFORE = "Improved codebase maintainability by 20% (peer review) via RESTful API design, CI/CD pipelines, and technical documentation across all prototypes."

# Calculate split positions
def find_second_by(text):
    """Find the second occurrence of ' by ' (not the one in ' by X%') or similar"""
    first_by = text.find(' by ', 5)
    if first_by >= 0:
        second_by = text.find(' by ', first_by + 4)
        return second_by
    return -1

def find_and_split(text):
    """Find ' and ' at a reasonable position"""
    and_pos = text.find(' and ', int(len(text) * 0.4))
    if and_pos > 0 and and_pos < len(text) - 20:
        return and_pos
    return -1

def find_via_split(text):
    via_pos = text.find(' via ', int(len(text) * 0.3))
    if via_pos > 0 and via_pos < len(text) - 15:
        return via_pos
    return -1

def find_with_split(text):
    with_pos = text.find(' with ', int(len(text) * 0.4))
    if with_pos > 0 and with_pos < len(text) - 15:
        return with_pos
    return -1

# P18: split at second "by"
pos = find_second_by(P18_TEXT)
if pos > 0:
    BULLET_SPLITS[18] = pos + 6  # after " by " (6 chars)

# P19: split at second "by"
pos = find_second_by(P19_TEXT)
if pos > 0:
    BULLET_SPLITS[19] = pos + 6

# P20: split at second "by" or "via"
pos = find_second_by(P20_TEXT)
if pos > 0:
    BULLET_SPLITS[20] = pos + 6

# P22: split at "and" (two methods)
# "designing... and deploying..."
pos = P22_TEXT.find(' and ', 80)  # find the 'and' joining the two methods
if pos > 0:
    BULLET_SPLITS[22] = pos + 5  # after " and "

# P23: split at second "by" or "with"
pos = find_second_by(P23_TEXT)
if pos > 0:
    BULLET_SPLITS[23] = pos + 6

# P25: split at "with"
pos = P25_TEXT.find(' with ', 70)
if pos > 0:
    BULLET_SPLITS[25] = pos + 6  # after " with "

# P26: split at "and" (two parallel actions)
pos = P26_TEXT.find(' and ', 60)
if pos > 0:
    BULLET_SPLITS[26] = pos + 5  # after " and "

# P27: split at "via" or after a comma
# "via RESTful API design, CI/CD pipelines, and technical documentation..."
pos = P27_TEXT_BEFORE.find(' across all prototypes')
if pos > 0:
    BULLET_SPLITS[27] = pos
else:
    pos = P27_TEXT_BEFORE.find(' via ', 30)
    if pos > 0:
        BULLET_SPLITS[27] = pos + 5

print("Splits defined:")
for idx, sp in sorted(BULLET_SPLITS.items()):
    full = doc.paragraphs[idx].text
    bullet = get_bullet_char(full)
    text = full[len(bullet):].lstrip() if bullet else full
    print(f'P{idx}: at {sp}')
    print(f'  1: {text[:sp- (1 if bullet else 0)].strip()}')
    print(f'  2: {text[sp- (1 if bullet else 0):].strip()}')

# Actually, let me recompute splits on the actual paragraph text
BULLET_SPLITS = {}
for i in [18, 19, 20, 22, 23, 25, 26, 27]:
    p = doc.paragraphs[i]
    full = p_text(p)
    bullet = get_bullet_char(full)
    text = full[len(bullet):].lstrip() if bullet else full
    
    if i == 18:
        # Split at second "by"
        pos = text.find(' by ', text.find(' by ') + 4)
        if pos > 0:
            BULLET_SPLITS[i] = pos + len(' by ') + len(bullet) + 1  # account for bullet prefix
    
    # ... etc

# Actually let me just manually compute based on string content
# I'll read the exact text and find the split points

for i in [18, 19, 20, 22, 23, 25, 26, 27]:
    p = doc.paragraphs[i]
    full = p_text(p)
    print(f'\nP{i} ({len(full)} chars): {repr(full)}')
