"""Rebuild v3 with v2 spacing, NO bullet splitting, minimized spacers for 1-page"""
from docx import Document
from lxml import etree
import copy

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
doc = Document(r'C:\Users\786 COMPUTERS\Downloads\resume\Resume_Syed_Abdullah_Yaqoob_v3.docx')

def get_all_t(p_elem):
    """Get all w:t elements in document order"""
    return list(p_elem.iter(f'{{{W_NS}}}t'))

def is_empty_para(p):
    """Check if paragraph is an empty separator"""
    for t in get_all_t(p._element):
        if t.text and t.text.strip():
            return False
    return True

def is_thin_spacer(p):
    """Check if paragraph is a thin spacer (whitespace only)"""
    t = p.text
    return t is not None and t.strip() == '' and len(t) <= 2

def set_spacing(p_elem, line, rule, after='0'):
    ppr = p_elem.find(f'{{{W_NS}}}pPr', None)
    if ppr is None:
        ppr = etree.SubElement(p_elem, f'{{{W_NS}}}pPr')
    spacing = ppr.find(f'{{{W_NS}}}spacing', None)
    if spacing is None:
        spacing = etree.SubElement(ppr, f'{{{W_NS}}}spacing')
    spacing.set(f'{{{W_NS}}}line', str(line))
    spacing.set(f'{{{W_NS}}}lineRule', rule)
    spacing.set(f'{{{W_NS}}}after', after)

# Fix ALL paragraphs
for i, p in enumerate(doc.paragraphs):
    if is_empty_para(p):
        # Section separator - minimize
        set_spacing(p._element, 40, 'exact')  # ~2pt instead of 4pt
    elif is_thin_spacer(p):
        # Thin spacer between sub-sections - nearly remove
        set_spacing(p._element, 1, 'exact')  # ~0.05pt - barely visible
    else:
        # Body text - v2 style
        set_spacing(p._element, 200, 'auto')

# Save
out = r'C:\Users\786 COMPUTERS\Downloads\resume\Resume_Syed_Abdullah_Yaqoob_v3_fixed.docx'
doc.save(out)
print(f'Saved: {out}')

# Verify
doc2 = Document(out)
print(f'Paragraphs: {len(doc2.paragraphs)}')
for i, p in enumerate(doc2.paragraphs):
    txt = p.text[:100] if p.text else '(empty)'
    print(f'P{i}: {repr(txt)}')
