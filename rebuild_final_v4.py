"""v3: v2 spacing, clean bullet splits with ALL words preserved, 1-page"""
from docx import Document
from lxml import etree
import copy

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
doc = Document(r'C:\Users\786 COMPUTERS\Downloads\resume\Resume_Syed_Abdullah_Yaqoob_v3.docx')

def p_text(p):
    return ''.join(t.text or '' for t in p._element.iter(f'{{{W_NS}}}t'))

def clear_text(p_elem):
    for t in p_elem.iter(f'{{{W_NS}}}t'):
        t.text = ''

def set_text(p_elem, text):
    first_t = p_elem.find(f'.//{{{W_NS}}}t', None)
    if first_t is not None:
        first_t.text = text
        first_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    for t in list(p_elem.iter(f'{{{W_NS}}}t'))[1:]:
        t.text = ''

def set_spacing(p_elem, line, rule, after='0'):
    ppr = p_elem.find(f'{{{W_NS}}}pPr', None)
    if ppr is None: ppr = etree.SubElement(p_elem, f'{{{W_NS}}}pPr')
    spacing = ppr.find(f'{{{W_NS}}}spacing', None)
    if spacing is None: spacing = etree.SubElement(ppr, f'{{{W_NS}}}spacing')
    spacing.set(f'{{{W_NS}}}line', str(line))
    spacing.set(f'{{{W_NS}}}lineRule', rule)
    spacing.set(f'{{{W_NS}}}after', after)

def get_bullet_char(text):
    for c in '\u2022\u25CF\uFFFD\u2023\u2219':
        if text.startswith(c): return c
    return ''

# Step 1: Fix spacing
for p in doc.paragraphs:
    t = p.text
    if t == '':
        set_spacing(p._element, 80, 'exact')
    elif t is not None and t.strip() == '' and len(t) <= 2:
        set_spacing(p._element, 1, 'exact')
    else:
        set_spacing(p._element, 200, 'auto')

# Step 2: Define clean splits
# Each split point is where the SECOND part starts in the original text
split_texts = {
    18: "deployed on AWS Lambda and S3 for scalable serverless inference.",
    19: "by building RESTful Node.js/Express microservices on AWS EC2 with Docker containerization.",
    20: "by automating CI/CD pipelines via GitHub Actions with Kubernetes orchestration for zero-downtime rollouts.",
    22: "by designing a full-year Agile/Scrum roadmap of 6+ events/semester and deploying the society site on AWS S3 + Vercel.",
    23: "by automating infrastructure via CI/CD GitHub Actions workflows with Salesforce CRM for member engagement tracking.",
    25: "by developing modular Python/Node.js components with TensorFlow inference and Pandas/NumPy data pipelines.",
    26: "by containerizing services with Docker and orchestrating via Kubernetes on AWS EC2.",
}

splits = {}
for idx, second in split_texts.items():
    full = p_text(doc.paragraphs[idx])
    pos = full.find(second)
    if pos > 0:
        splits[idx] = pos

# Handle P27 (has PROJECTS)
p27 = doc.paragraphs[27]
full27 = p_text(p27)
proj_pos = full27.find('\n\n')
if proj_pos > 0:
    bullet_part = full27[:proj_pos]
    via_pos = bullet_part.find(' via ')
    if via_pos > 0:
        splits[27] = via_pos

# Apply splits in reverse
for idx in sorted(splits.keys(), reverse=True):
    p = doc.paragraphs[idx]
    full = p_text(p)
    sp = splits[idx]
    bullet = get_bullet_char(full)
    
    if idx == 27:
        proj_idx = full.find('\n\n')
        if proj_idx > 0:
            bullet_text = full[:proj_idx]
        else:
            bullet_text = full
        
        if sp < len(bullet_text):
            first = bullet_text[:sp].rstrip()
            second = bullet_text[sp:].strip()
            
            clear_text(p._element)
            set_text(p._element, first)
            
            new_p = copy.deepcopy(p._element)
            p._element.addnext(new_p)
            clear_text(new_p)
            set_text(new_p, f'{bullet} {second}')
            set_spacing(new_p, 200, 'auto')
            
            proj_heading = copy.deepcopy(p._element)
            new_p.addnext(proj_heading)
            clear_text(proj_heading)
            set_text(proj_heading, 'PROJECTS')
            for r in proj_heading.iter(f'{{{W_NS}}}r'):
                rpr = r.find(f'{{{W_NS}}}rPr', None)
                if rpr is None: rpr = etree.SubElement(r, f'{{{W_NS}}}rPr')
                b = rpr.find(f'{{{W_NS}}}b', None)
                if b is None: b = etree.SubElement(rpr, f'{{{W_NS}}}b')
            set_spacing(proj_heading, 200, 'auto')
        continue
    
    first = full[:sp].rstrip()
    second = full[sp:].strip()
    if not second:
        continue
    
    clear_text(p._element)
    set_text(p._element, first)
    
    new_p = copy.deepcopy(p._element)
    p._element.addnext(new_p)
    clear_text(new_p)
    set_text(new_p, f'{bullet} {second}')
    set_spacing(new_p, 200, 'auto')

out = r'C:\Users\786 COMPUTERS\Downloads\resume\Resume_Syed_Abdullah_Yaqoob_v3_fixed.docx'
doc.save(out)
print(f'Saved: {out}')

doc2 = Document(out)
print(f'Paragraphs: {len(doc2.paragraphs)}')
for i, p in enumerate(doc2.paragraphs):
    txt = p.text
    if not txt: continue
    if any(c in txt for c in '\u2022\u25CF\uFFFD') and len(txt) > 20:
        print(f'P{i} ({len(txt)}c): {txt[:120]}')
    elif 'PROJECTS' in txt:
        print(f'P{i}: {repr(txt[:60])}')
