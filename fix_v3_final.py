"""
Fix Resume v3 to match v2 style:
1. Change line spacing from 184 (0.767) to 200 (0.833) for body text
2. Change separator lines from 50 to 80 exact
3. Change thin spacers from 50 to 80 where needed
4. Split long employment bullets into two (without removing any words)
5. Handle PROJECTS heading embedded in last NRAI bullet
"""

from docx import Document
from lxml import etree
import copy

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def get_all_runs(p_elem):
    """Get all runs including those inside hyperlinks"""
    runs = list(p_elem.findall(f'{{{W_NS}}}r', None))
    for hl in p_elem.findall(f'{{{W_NS}}}hyperlink', None):
        runs.extend(hl.findall(f'{{{W_NS}}}r', None))
    return runs

def clear_text(p_elem):
    for r in get_all_runs(p_elem):
        t = r.find(f'{{{W_NS}}}t', None)
        if t is not None:
            t.text = ''

def set_text_first_run(p_elem, text):
    runs = get_all_runs(p_elem)
    if not runs:
        r = etree.SubElement(p_elem, f'{{{W_NS}}}r')
        t = etree.SubElement(r, f'{{{W_NS}}}t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text
        return
    for j, r in enumerate(runs):
        t_elem = r.find(f'{{{W_NS}}}t', None)
        if t_elem is not None:
            t_elem.text = text if j == 0 else ''

def get_full_text(p):
    parts = []
    for r in get_all_runs(p._element):
        t = r.find(f'{{{W_NS}}}t', None)
        if t is not None and t.text:
            parts.append(t.text)
    return ''.join(parts)

def has_hyperlink(p):
    return len(p._element.findall(f'{{{W_NS}}}hyperlink', None)) > 0

def set_spacing(ppr, line_val, rule, after='0'):
    spacing = ppr.find(f'{{{W_NS}}}spacing', None)
    if spacing is None:
        spacing = etree.SubElement(ppr, f'{{{W_NS}}}spacing')
    spacing.set(f'{{{W_NS}}}line', str(line_val))
    spacing.set(f'{{{W_NS}}}lineRule', rule)
    spacing.set(f'{{{W_NS}}}after', after)

# ===== MAIN =====
doc = Document(r'C:\Users\786 COMPUTERS\Downloads\resume\Resume_Syed_Abdullah_Yaqoob_v3.docx')

# STEP 1: Fix spacing on ALL paragraphs
for i, p in enumerate(doc.paragraphs):
    ppr = p._element.find(f'{{{W_NS}}}pPr', None)
    if ppr is None:
        ppr = etree.SubElement(p._element, f'{{{W_NS}}}pPr')
    
    text = p.text
    
    if text == '':
        runs = get_all_runs(p._element)
        is_empty = all(r.find(f'{{{W_NS}}}t', None) is None or 
                      (r.find(f'{{{W_NS}}}t', None).text or '').strip() == '' 
                      for r in runs)
        if is_empty:
            set_spacing(ppr, 80, 'exact')  # Separator
        else:
            set_spacing(ppr, 200, 'auto')  # Regular body
    elif text.strip() == '' and len(text) <= 2:
        if i in (5, 14, 59):
            set_spacing(ppr, 120, 'auto')
        elif i in (64, 68, 71):
            set_spacing(ppr, 80, 'auto')
        else:
            set_spacing(ppr, 120, 'auto')
    else:
        set_spacing(ppr, 200, 'auto')

# STEP 2: Split long employment bullet points
# Employment bullets: P18, P19, P20, P22, P23, P25, P26, P27
# P27 also contains embedded PROJECTS heading
# None of these have hyperlinks - safe to modify text

def find_split(text):
    if len(text) <= 130:
        return None
    mid = len(text) // 2
    best = None
    best_d = float('inf')
    for phrase in [' by ', ' with ', ' using ', ' for ', ' via ', ' across ',
                   ' through ', ' serving ', ' deploying ', ' and ', ' of ', ' in ']:
        s = int(len(text) * 0.35)
        while True:
            pos = text.find(phrase, s)
            if pos == -1 or pos > len(text) - 15:
                break
            d = abs(pos - mid)
            if d < best_d:
                best_d = d
                best = pos + len(phrase)
            s = pos + 1
    if best and best < len(text) - 10:
        return best
    return None

# Collect splits
splits = {}
for i in [18, 19, 20, 22, 23, 25, 26, 27]:
    full = get_full_text(doc.paragraphs[i])
    # For P27, only consider the part before \n\n (the PROJECTS separator)
    if i == 27:
        idx = full.find('\n\n')
        if idx > 0:
            text = full[:idx]
        else:
            text = full
    else:
        text = full
    
    sp = find_split(text)
    if sp:
        splits[i] = (sp, full)  # Store split point and full text

# Apply splits in reverse
for idx in sorted(splits.keys(), reverse=True):
    sp, full = splits[idx]
    p = doc.paragraphs[idx]
    
    # Handle P27 specially (has PROJECTS embedded)
    if idx == 27:
        # Find where PROJECTS part starts
        proj_idx = full.find('\n\nPROJECTS')
        if proj_idx > 0:
            bullet_text = full[:proj_idx]
            projects_part = full[proj_idx:]
            # Split only the bullet part
            if sp < proj_idx:
                first = bullet_text[:sp].rstrip()
                second = bullet_text[sp:].strip()
                
                # Get bullet char
                bullet = ''
                for c in ['\u2022', '\u25CF', '\uFFFD', '\u2023', '\u2219']:
                    if full.startswith(c):
                        bullet = c
                        break
                
                # First part stays in original paragraph
                clear_text(p._element)
                set_text_first_run(p._element, first)
                
                # Create new paragraph for second bullet part
                new_p = copy.deepcopy(p._element)
                p._element.addnext(new_p)
                clear_text(new_p)
                set_text_first_run(new_p, f'{bullet} {second}')
                # Fix spacing on new para
                nppr = new_p.find(f'{{{W_NS}}}pPr', None)
                if nppr is None:
                    nppr = etree.SubElement(new_p, f'{{{W_NS}}}pPr')
                set_spacing(nppr, 200, 'auto')
                
                # Now the PROJECTS part needs to be moved after the separator
                # Actually, let's keep it simple - just don't separate PROJECTS
                # Instead, create the PROJECTS paragraph after the separator
                # We need to find the separator (original P28, now shifted)
                # This is complex, so let's just add PROJECTS as a new paragraph
                # after the separator
        continue
    
    # Regular bullet splitting
    first_part = full[:sp].rstrip()
    second_part = full[sp:].strip()
    if not second_part:
        continue
    
    bullet = ''
    for c in ['\u2022', '\u25CF', '\uFFFD', '\u2023', '\u2219']:
        if full.startswith(c):
            bullet = c
            break
    
    new_p = copy.deepcopy(p._element)
    p._element.addnext(new_p)
    
    clear_text(p._element)
    set_text_first_run(p._element, first_part)
    
    clear_text(new_p)
    set_text_first_run(new_p, f'{bullet} {second_part}')
    
    nppr = new_p.find(f'{{{W_NS}}}pPr', None)
    if nppr is None:
        nppr = etree.SubElement(new_p, f'{{{W_NS}}}pPr')
    set_spacing(nppr, 200, 'auto')

# STEP 3: Handle P27 (last NRAI bullet with PROJECTS)
# Find the paragraph containing PROJECTS heading
for i, p in enumerate(doc.paragraphs):
    full = get_full_text(p)
    if '\n\nPROJECTS' in full:
        # This paragraph has the embedded PROJECTS heading
        # Extract PROJECTS as its own paragraph
        bullet_text, projects_part = full.split('\n\n', 1)
        if projects_part.strip() == 'PROJECTS':
            clear_text(p._element)
            set_text_first_run(p._element, bullet_text)
            
            # Create PROJECTS paragraph after this one
            new_p = copy.deepcopy(p._element)
            p._element.addnext(new_p)
            
            # Need to handle the newlines - they were in separate runs in the original
            # Just set the text
            clear_text(new_p)
            
            # Make it bold by copying the bold run style from the original PROJECTS run
            # Find the original PROJECTS run to copy its bold formatting
            set_text_first_run(new_p, 'PROJECTS')
            
            # Make it bold
            runs = new_p.findall(f'{{{W_NS}}}r', None)
            if runs:
                rpr = runs[0].find(f'{{{W_NS}}}rPr', None)
                if rpr is None:
                    rpr = etree.SubElement(runs[0], f'{{{W_NS}}}rPr')
                b = rpr.find(f'{{{W_NS}}}b', None)
                if b is None:
                    b = etree.SubElement(rpr, f'{{{W_NS}}}b')
            
            # Fix spacing on PROJECTS paragraph
            nppr = new_p.find(f'{{{W_NS}}}pPr', None)
            if nppr is None:
                nppr = etree.SubElement(new_p, f'{{{W_NS}}}pPr')
            set_spacing(nppr, 200, 'auto')
        
        break

# STEP 4: Fix trailing newline in MiniChat bullet (P57 original)
for i, p in enumerate(doc.paragraphs):
    full = get_full_text(p)
    if full.rstrip('\n') == full.strip('\n') and full.endswith('\n') and full.startswith('\uFFFD'):
        # Trailing newline in a bullet
        clean = full.rstrip('\n')
        clear_text(p._element)
        set_text_first_run(p._element, clean)
        break

# Save
output_path = r'C:\Users\786 COMPUTERS\Downloads\resume\Resume_Syed_Abdullah_Yaqoob_v3_fixed.docx'
doc.save(output_path)
print(f'Saved: {output_path}')

# Verify
doc2 = Document(output_path)
print(f'Paragraphs: {len(doc2.paragraphs)}')
for i, p in enumerate(doc2.paragraphs):
    txt = get_full_text(p)
    if '\n' in txt:
        print(f'P{i}: {repr(txt[:80])}')
    elif txt:
        print(f'P{i}: {txt[:110]}')
    else:
        print(f'P{i}: (empty)')
