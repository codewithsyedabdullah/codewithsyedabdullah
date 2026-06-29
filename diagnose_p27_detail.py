"""Detailed P27 diagnosis - all elements and text"""
from docx import Document
from lxml import etree

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
doc = Document(r'C:\Users\786 COMPUTERS\Downloads\resume\Resume_Syed_Abdullah_Yaqoob_v3.docx')
p27 = doc.paragraphs[27]
elem = p27._element

# Print full XML of P27
xml_str = etree.tostring(elem, pretty_print=True, encoding='unicode')
print("=== P27 XML ===")
print(xml_str[:2000])
print()

# Also look at P17-P26 for the NRAI employment bullets to see if there are more that need splitting
print("=== NRAI bullets (P25-P27) ===")
for i in range(25, 28):
    p = doc.paragraphs[i]
    full = ''
    for t in p._element.iter(f'{{{W_NS}}}t'):
        full += t.text or ''
    print(f'P{i} ({len(full)} chars): {full[:200]}')
    brs = p._element.findall(f'.//{{{W_NS}}}br')
    print(f'  w:br count: {len(brs)}')
    ts = list(p._element.iter(f'{{{W_NS}}}t'))
    print(f'  t elements: {len(ts)}')
    for j, t in enumerate(ts):
        print(f'    t[{j}]: len={len(t.text or "")} text={repr((t.text or "")[:150])}')
    print()
