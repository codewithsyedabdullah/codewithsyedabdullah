"""Check output for leftover w:br in P35 and structure around PROJECTS"""
from docx import Document
from lxml import etree

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
doc = Document(r'C:\Users\786 COMPUTERS\Downloads\resume\Resume_Syed_Abdullah_Yaqoob_v3_fixed.docx')

# Check P34-P39 structure
print("=== P34-P39 ===")
for i in range(34, 40):
    p = doc.paragraphs[i]
    txt = p.text
    elem = p._element
    brs = list(elem.iter(f'{{{W_NS}}}br'))
    ts = list(elem.iter(f'{{{W_NS}}}t'))
    spacing = elem.find(f'.//{{{W_NS}}}spacing', None)
    line = spacing.get(f'{{{W_NS}}}line') if spacing is not None else 'N/A'
    rule = spacing.get(f'{{{W_NS}}}lineRule') if spacing is not None else 'N/A'
    print(f'P{i}: txt={repr(txt[:60]) if txt else "(empty)"} br={len(brs)} t={len(ts)} line={line}/{rule}')

# Check PROJECTS run formatting
print()
print("=== P36 (PROJECTS) runs ===")
p36 = doc.paragraphs[36]
for r in p36._element.iter(f'{{{W_NS}}}r'):
    rpr = r.find(f'{{{W_NS}}}rPr', None)
    b = rpr.find(f'{{{W_NS}}}b', None) if rpr is not None else None
    t = r.find(f'{{{W_NS}}}t', None)
    print(f'  t={repr(t.text) if t is not None else "None"} bold={b is not None}')
